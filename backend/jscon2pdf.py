from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from matplotlib import pyplot as plt
from matplotlib import rc
import os
import uuid

def add_page_number(paragraph):
    """Добавляет номер страницы в нижний колонтитул."""
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)

    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run._r.append(instrText)

    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def create_element(name):
    """Создает XML элемент."""
    return OxmlElement(name)

def create_attribute(element, name, value):
    """Создает атрибут для XML элемента."""
    element.set(qn(name), value)

def add_page_border(document):
    """Добавляет рамку страницы."""
    sectPr = document.sections[0]._sectPr
    pgBorders = create_element('w:pgBorders')
    create_attribute(pgBorders, 'w:offsetFrom', 'page')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = create_element(f'w:{border_name}')
        create_attribute(border, 'w:val', 'single')
        create_attribute(border, 'w:sz', '4')
        create_attribute(border, 'w:space', '24')
        create_attribute(border, 'w:color', '000000')
        pgBorders.append(border)
    
    sectPr.append(pgBorders)

def latex_to_image(latex_code, output_path):
    """Преобразует LaTeX-формулу в изображение с улучшенным стилем."""
    rc('text', usetex=False)
    fig = plt.figure(figsize=(10, 3))
    fig.patch.set_facecolor('white')
    fig.text(0.5, 0.5, f"${latex_code}$", ha='center', va='center', fontsize=24, 
             color='navy')
    plt.axis('off')
    fig.savefig(output_path, dpi=300, bbox_inches='tight', pad_inches=0.3, 
                facecolor='white')
    plt.close(fig)

def create_custom_styles(document):
    """Создает пользовательские стили документа."""
    # Стиль заголовка документа
    style = document.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(24)
    font.bold = True
    font.color.rgb = RGBColor(0, 32, 96)
    
    # Стиль заголовка формулы
    style = document.styles.add_style('FormulaTitle', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(16)
    font.bold = True
    font.color.rgb = RGBColor(47, 84, 150)
    
    # Стиль описания
    style = document.styles.add_style('Description', WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

def json_to_docx(json_data, output_directory):
    """Преобразует JSON-данные в стильный docx-файл."""
    file_name = f"mathematical_formulas_{uuid.uuid4().hex[:8]}.docx"
    output_file = os.path.join(output_directory, file_name)

    if not os.path.exists(output_directory):
        os.makedirs(output_directory)

    document = Document()
    create_custom_styles(document)
    add_page_border(document)

    # Настройка полей страницы
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Добавление заголовка
    title = document.add_paragraph("Математические формулы", style='CustomTitle')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph()  # Пустая строка после заголовка

    for i, item in enumerate(json_data, 1):
        formula = item.get("latex_formula", "")
        legend = item.get("legend", f"Формула {i}")
        description = item.get("description", "")

        # Добавление номера и легенды формулы
        formula_title = document.add_paragraph(style='FormulaTitle')
        formula_title.add_run(f"Формула {i}: {legend}")
        formula_title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Добавление формулы как изображения
        image_path = f"formula_{uuid.uuid4().hex}.png"
        latex_to_image(formula, image_path)
        
        formula_paragraph = document.add_paragraph()
        formula_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        formula_run = formula_paragraph.add_run()
        formula_run.add_picture(image_path, width=Inches(6))
        os.remove(image_path)

        # Добавление описания
        if description:
            desc_paragraph = document.add_paragraph(style='Description')
            desc_paragraph.add_run(description)
            desc_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        # Добавление разделителя (кроме последней формулы)
        if i < len(json_data):
            document.add_paragraph()
            separator = document.add_paragraph()
            separator.alignment = WD_ALIGN_PARAGRAPH.CENTER
            separator.add_run('⚡').font.size = Pt(14)
            document.add_paragraph()
        
        # Добавление номера страницы
        footer = document.sections[0].footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer_paragraph)

    document.save(output_file)
    return output_file

# Пример использования
if __name__ == "__main__":
    json_data = [
        {
            "latex_formula": "E = mc^2",
            "legend": "Формула энергии",
            "description": "Это знаменитая формула Эйнштейна, показывающая взаимосвязь между массой и энергией. Она является одним из самых известных уравнений в физике и демонстрирует, что энергия и масса являются разными формами одного и того же."
        },
        {
            "latex_formula": "F = ma",
            "legend": "Второй закон Ньютона",
            "description": "Фундаментальный закон механики, описывающий взаимосвязь между силой, действующей на тело, его массой и ускорением. Этот закон лежит в основе классической механики."
        }
    ]

    output_dir = "output_docs"
    output_file = json_to_docx(json_data, output_dir)
    print(f"Документ успешно сохранен: {output_file}")